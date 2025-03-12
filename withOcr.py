import fitz  
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from pdf2docx import Converter

# coloca o caminho dos arquivos aqui
pdf_path = "documento.pdf"
word_path = "documento.docx"

# Verifica se o PDF tem texto selecionável
doc_pdf = fitz.open(pdf_path)
has_text = any(page.get_text() for page in doc_pdf)

if has_text:
    # Se o PDF já contém texto, usa pdf2docx
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()
else:
    # Se o PDF é escaneado, usa OCR para extrair texto
    doc = Document()
    images = convert_from_path(pdf_path)
    
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img, lang="por")
        doc.add_paragraph(f"Página {i+1}")
        doc.add_paragraph(text)
        doc.add_paragraph("\n---\n")  # Separador entre páginas
    
    doc.save(word_path)

print(f"Conversão concluída! Arquivo salvo como {word_path}")
